from fastapi import APIRouter, File, UploadFile, Form, HTTPException
from fastapi.responses import JSONResponse
from core.config import STATIC_DIR, MAX_UPLOAD_SIZE, ALLOWED_EXTENSIONS
from core.utils import is_allowed_file, generate_temp_filename, cleanup_temp_file, logger
from models.ocr import extract_text
from models.translator import translator_model
import shutil
import easyocr
from PIL import Image
import numpy as np
import io
from typing import Optional
import PyPDF2
from docx import Document

router = APIRouter(prefix="/ocr", tags=["OCR"])

# Initialize EasyOCR reader (lazy loading)
ocr_reader = None

def get_ocr_reader():
    """Lazy load OCR reader"""
    global ocr_reader
    if ocr_reader is None:
        try:
            # Initialize with Nepali and English
            ocr_reader = easyocr.Reader(['ne', 'en'], gpu=True)
            logger.info("EasyOCR reader initialized successfully")
        except Exception as e:
            logger.error(f"Failed to initialize OCR with GPU: {e}")
            # Fallback to CPU
            ocr_reader = easyocr.Reader(['ne', 'en'], gpu=False)
    return ocr_reader

def detect_language(text: str) -> str:
    """Simple language detection based on Unicode ranges"""
    # Devanagari script (Nepali/Hindi)
    if any('\u0900' <= char <= '\u097F' for char in text):
        return 'nepali'
    # Sinhala script
    if any('\u0D80' <= char <= '\u0DFF' for char in text):
        return 'sinhala'
    # Tamil script
    if any('\u0B80' <= char <= '\u0BFF' for char in text):
        return 'tamil'
    # Bengali script
    if any('\u0980' <= char <= '\u09FF' for char in text):
        return 'bengali'
    # Default to English
    return 'english'


@router.post("/")
async def ocr_image(file: UploadFile = File(...), detail_level: int = Form(default=1)):
    """
    Extract text from uploaded image using OCR
    
    Args:
        file: Image file (JPG, PNG, etc.)
        detail_level: 0 = simple, 1 = detailed (with confidence scores)
    
    Returns:
        JSON with extracted text and metadata
    """
    try:
        # Read image file
        contents = await file.read()
        image = Image.open(io.BytesIO(contents))
        
        # Convert to numpy array for EasyOCR
        image_np = np.array(image)
        
        # Perform OCR
        reader = get_ocr_reader()
        results = reader.readtext(image_np, detail=detail_level)
        
        # Extract text
        if detail_level == 0:
            extracted_text = ' '.join(results)
        else:
            extracted_text = ' '.join([text for (bbox, text, conf) in results])
            
        if not extracted_text.strip():
            raise HTTPException(status_code=400, detail="No text found in image")
        
        # Prepare response
        response = {
            "success": True,
            "extracted_text": extracted_text.strip(),
            "word_count": len(extracted_text.split()),
            "confidence": sum([conf for (_, _, conf) in results]) / len(results) if results else 0,
            "filename": file.filename
        }
        
        # Add detailed results if requested
        if detail_level == 1:
            response["details"] = [
                {
                    "text": text,
                    "confidence": float(conf),
                    "bbox": [[int(coord) for coord in point] for point in bbox]
                }
                for (bbox, text, conf) in results
            ]
        
        return JSONResponse(content=response)
        
    except Exception as e:
        logger.error(f"OCR error: {str(e)}")
        raise HTTPException(status_code=500, detail=f"OCR processing failed: {str(e)}")


@router.post("/translate/")
async def extract_and_translate(
    file: UploadFile = File(...),
    source_lang: str = Form(default="nepali"),
    target_lang: str = Form(default="english")
):
    """
    Extract text from image and translate it
    
    Args:
        file: Image file
        source_lang: Source language (nepali, sinhala, etc.)
        target_lang: Target language (english)
    
    Returns:
        JSON with extracted text and translation
    """
    try:
        # Read image file
        contents = await file.read()
        image = Image.open(io.BytesIO(contents))
        
        # Convert to numpy array
        image_np = np.array(image)
        
        # Perform OCR
        reader = get_ocr_reader()
        results = reader.readtext(image_np, detail=1)
        
        if not results:
            raise HTTPException(status_code=400, detail="No text found in image")
        
        # Extract text
        extracted_text = ' '.join([text for (bbox, text, conf) in results])
        avg_confidence = sum([conf for (_, _, conf) in results]) / len(results)
        
        # Translate the extracted text
        if translator_model.model is None:
            translator_model.load_model()
        
        translation = translator_model.translate(
            extracted_text,
            source_lang=source_lang,
            target_lang=target_lang
        )
        
        # Detect language
        detected_lang = detect_language(extracted_text)
        
        return JSONResponse(content={
            "success": True,
            "extracted_text": extracted_text.strip(),
            "translation": translation,
            "source_language": source_lang,
            "detected_language": detected_lang,
            "target_language": target_lang,
            "confidence": float(avg_confidence),
            "word_count": len(extracted_text.split()),
            "ocr_details": [
                {
                    "text": text,
                    "confidence": float(conf)
                }
                for (bbox, text, conf) in results
            ]
        })
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"OCR + Translation error: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Processing failed: {str(e)}")


@router.post("/live-translate/")
async def live_image_translation(
    file: UploadFile = File(...),
    source_lang: str = Form(default="nepali"),
    target_lang: str = Form(default="english"),
    return_image: bool = Form(default=False)
):
    """
    Live image translation with text region overlays (like Google Lens)
    
    Returns text regions with bounding boxes and translations
    """
    try:
        # Read image file
        contents = await file.read()
        image = Image.open(io.BytesIO(contents))
        image_np = np.array(image)
        
        # Perform OCR with bounding boxes
        reader = get_ocr_reader()
        results = reader.readtext(image_np, detail=1)
        
        if not results:
            raise HTTPException(status_code=400, detail="No text found in image")
        
        # Load translator if not already loaded
        if translator_model.model is None:
            translator_model.load_model()
        
        # Process each text region
        text_regions = []
        total_confidence = 0
        
        for (bbox, text, conf) in results:
            # Translate the text
            translation = translator_model.translate(
                text,
                source_lang=source_lang,
                target_lang=target_lang
            )
            
            text_regions.append({
                "bbox": [[int(coord) for coord in point] for point in bbox],
                "original_text": text,
                "translated_text": translation,
                "confidence": float(conf)
            })
            total_confidence += conf
        
        avg_confidence = total_confidence / len(results) if results else 0
        
        return JSONResponse(content={
            "success": True,
            "text_regions": text_regions,
            "total_regions": len(text_regions),
            "avg_confidence": float(avg_confidence),
            "source_language": source_lang,
            "target_language": target_lang,
            "message": f"Successfully translated {len(text_regions)} text regions"
        })
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Live translation error: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Live translation failed: {str(e)}")


@router.get("/health/")
async def ocr_health_check():
    """Check OCR system status"""
    try:
        reader = get_ocr_reader()
        return {
            "status": "healthy",
            "ocr_ready": reader is not None,
            "supported_languages": ['ne', 'en'],
            "gpu_enabled": reader.gpu if reader else False
        }
    except Exception as e:
        return {
            "status": "error",
            "error": str(e)
        }


@router.post("/text-file")
async def extract_from_text_file(file: UploadFile = File(...)):
    """
    Extract text from uploaded .txt file.
    """
    try:
        if not file.filename or not file.filename.endswith('.txt'):
            raise HTTPException(status_code=400, detail="Only .txt files allowed")
        
        # Read text content
        content = await file.read()
        text = content.decode('utf-8')
        
        return {
            "extracted_text": text,
            "filename": file.filename,
            "character_count": len(text),
            "status": "success"
        }
        
    except UnicodeDecodeError:
        raise HTTPException(status_code=400, detail="File encoding not supported. Use UTF-8")
    except Exception as e:
        logger.error(f"Text file extraction error: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/pdf")
async def extract_from_pdf(file: UploadFile = File(...)):
    """
    Extract text from uploaded PDF file.
    """
    try:
        if not file.filename or not file.filename.lower().endswith('.pdf'):
            raise HTTPException(status_code=400, detail="Only .pdf files allowed")
        
        # Read PDF content
        content = await file.read()
        pdf_reader = PyPDF2.PdfReader(io.BytesIO(content))
        
        # Extract text from all pages
        extracted_text = ""
        page_count = len(pdf_reader.pages)
        
        for page_num, page in enumerate(pdf_reader.pages):
            page_text = page.extract_text()
            extracted_text += f"\n--- Page {page_num + 1} ---\n{page_text}\n"
        
        return {
            "extracted_text": extracted_text.strip(),
            "filename": file.filename,
            "page_count": page_count,
            "character_count": len(extracted_text),
            "status": "success"
        }
        
    except Exception as e:
        logger.error(f"PDF extraction error: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to process PDF: {str(e)}")


@router.post("/docx")
async def extract_from_docx(file: UploadFile = File(...)):
    """
    Extract text from uploaded DOCX file.
    """
    try:
        if not file.filename or not (file.filename.lower().endswith('.docx') or file.filename.lower().endswith('.doc')):
            raise HTTPException(status_code=400, detail="Only .docx/.doc files allowed")
        
        # Read DOCX content
        content = await file.read()
        doc = Document(io.BytesIO(content))
        
        # Extract text from all paragraphs
        extracted_text = ""
        paragraph_count = 0
        
        for para in doc.paragraphs:
            if para.text.strip():
                extracted_text += para.text + "\n"
                paragraph_count += 1
        
        # Extract text from tables
        table_count = len(doc.tables)
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join([cell.text for cell in row.cells])
                extracted_text += row_text + "\n"
        
        return {
            "extracted_text": extracted_text.strip(),
            "filename": file.filename,
            "paragraph_count": paragraph_count,
            "table_count": table_count,
            "character_count": len(extracted_text),
            "status": "success"
        }
        
    except Exception as e:
        logger.error(f"DOCX extraction error: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to process DOCX: {str(e)}")
